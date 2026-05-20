"""
Summit AI Visibility Audit Tool v2.1
Powered by Google Gemini 2.5 Flash
"""

import streamlit as st
import requests
from bs4 import BeautifulSoup
import json
import re
import base64
import io
import os
import time
from datetime import datetime
from urllib.parse import urlparse

st.set_page_config(
    page_title="Summit AI Visibility Audit",
    page_icon="🔍",
    layout="wide",
    initial_sidebar_state="expanded"
)

import plotly.graph_objects as go
import google.genai as genai
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from fpdf import FPDF
from PIL import Image

# =============================================================================
# BRAND PALETTE
# =============================================================================
S_RED       = "#D0380B"
S_SIDEBAR   = "#262626"   # Dark grey sidebar as requested
S_CHARCOAL  = "#3E405B"   # Used for headings / table headers
S_WHITE     = "#FFFFFF"
S_OFFWHITE  = "#F7F7F8"
S_LIGHT     = "#EDEDF0"
S_INK       = "#1C1C2E"
S_MUTED     = "#6B6B80"
S_CAPTION   = "#9999AA"

BG_RED   = "#FDE9E5";  FG_RED   = "#B52D0A"
BG_AMBER = "#FFF4E0";  FG_AMBER = "#8A4800"
BG_GREEN = "#E6F5EC";  FG_GREEN = "#1E6E3C"

S_RED_RGB      = (208, 56, 11)
S_CHARCOAL_RGB = (62, 64, 91)
S_SIDEBAR_RGB  = (38, 38, 38)
FG_RED_RGB     = (181, 45, 10)
FG_AMBER_RGB   = (138, 72, 0)
FG_GREEN_RGB   = (30, 110, 60)
BG_RED_HEX     = "FDE9E5"
BG_AMBER_HEX   = "FFF4E0"
BG_GREEN_HEX   = "E6F5EC"

DIMENSION_CONFIG = {
    "Crawlability & Bot Access":  {"weight":0.20,"icon":"🤖","color":S_CHARCOAL, "description":"Can AI crawlers access and parse the page?"},
    "Structured Data / Schema":   {"weight":0.18,"icon":"🏷️", "color":"#7C3AED","description":"Schema.org JSON-LD quality and completeness"},
    "LLM Content Signals":        {"weight":0.15,"icon":"🧠","color":S_RED,     "description":"Content clarity, factual density, E-E-A-T signals"},
    "Meta & SEO Signals":         {"weight":0.12,"icon":"🔍","color":"#0E7C4A","description":"Title, meta description, canonical, Open Graph"},
    "Heading Structure":          {"weight":0.10,"icon":"📋","color":"#0369A1","description":"Semantic H1–H6 hierarchy and content flow"},
    "ARIA Implementation":        {"weight":0.10,"icon":"♿","color":"#BE185D","description":"Accessibility attributes aiding AI content parsing"},
    "Link Quality":               {"weight":0.08,"icon":"🔗","color":"#78350F","description":"Anchor text quality, internal/external link health"},
    "Image Alt Text":             {"weight":0.07,"icon":"🖼️", "color":"#374151","description":"Alt text coverage, quality and keyword relevance"},
    "AI Search Health":           {"weight":0.05,"icon":"📡","color":"#B45309","description":"llms.txt presence, AI bot directives in robots.txt"},
    "Duplicate Content & Tags":   {"weight":0.05,"icon":"📄","color":"#6B6B80","description":"Canonical consistency, duplicate titles, thin content"},
}

LOGO_PATH = os.path.join(os.path.dirname(__file__), "summit_logo.png")

def load_logo_bytes() -> bytes:
    with open(LOGO_PATH, "rb") as f:
        return f.read()

def logo_b64(b: bytes) -> str:
    return base64.b64encode(b).decode()

# =============================================================================
# CSS  — fixes white-on-white with explicit dark text colours everywhere
# =============================================================================
def inject_css():
    st.markdown(f"""
    <style>
    @import url('https://fonts.googleapis.com/css2?family=Inter:wght@400;500;600;700;800&display=swap');
    html, body, [class*="css"] {{ font-family:'Inter',sans-serif; }}
    .main {{ background:{S_OFFWHITE}; }}

    /* ── Sidebar ── dark grey #262626 */
    [data-testid="stSidebar"] {{ background:{S_SIDEBAR} !important; }}
    [data-testid="stSidebar"] p,
    [data-testid="stSidebar"] span,
    [data-testid="stSidebar"] label,
    [data-testid="stSidebar"] div {{ color:rgba(255,255,255,0.88) !important; }}
    [data-testid="stSidebar"] h1,
    [data-testid="stSidebar"] h2,
    [data-testid="stSidebar"] h3 {{ color:#FFFFFF !important; }}
    /* Text inputs — broad selectors to survive Streamlit version changes */
    [data-testid="stSidebar"] input,
    [data-testid="stSidebar"] textarea,
    [data-testid="stSidebar"] .stTextInput input,
    [data-testid="stSidebar"] .stTextArea textarea,
    [data-testid="stSidebar"] [data-baseweb="input"] input,
    [data-testid="stSidebar"] [data-baseweb="textarea"] textarea {{
        background:rgba(255,255,255,0.10) !important;
        color:#FFFFFF !important;
        border:1px solid rgba(255,255,255,0.28) !important;
        border-radius:6px !important;
        caret-color:#FFFFFF !important;
    }}
    [data-testid="stSidebar"] input::placeholder,
    [data-testid="stSidebar"] textarea::placeholder,
    [data-testid="stSidebar"] .stTextInput input::placeholder,
    [data-testid="stSidebar"] .stTextArea textarea::placeholder {{
        color:rgba(255,255,255,0.42) !important;
    }}
    [data-testid="stSidebar"] [data-baseweb="base-input"],
    [data-testid="stSidebar"] [data-baseweb="input"],
    [data-testid="stSidebar"] [data-baseweb="textarea"] {{
        background:rgba(255,255,255,0.10) !important;
        border-color:rgba(255,255,255,0.28) !important;
    }}
    /* URL tab buttons inside sidebar */
    [data-testid="stSidebar"] .stTabs [data-baseweb="tab"] {{
        color:rgba(255,255,255,0.70) !important;
        background:transparent !important;
    }}
    [data-testid="stSidebar"] .stTabs [aria-selected="true"] {{
        color:#FFFFFF !important;
        border-bottom:2px solid {S_RED} !important;
    }}
    [data-testid="stSidebar"] .stTabs [data-baseweb="tab-list"] {{
        background:transparent !important;
        border-bottom:1px solid rgba(255,255,255,0.12) !important;
    }}
    /* Run button */
    [data-testid="stSidebar"] .stButton > button {{
        background:{S_RED} !important;
        color:#FFFFFF !important;
        font-weight:700 !important;
        border:none !important;
        border-radius:6px !important;
        width:100%;
        padding:0.65rem 1rem;
        font-size:0.95rem;
        letter-spacing:0.02em;
        transition:background 0.15s;
    }}
    [data-testid="stSidebar"] .stButton > button:hover {{ background:#B52D0A !important; }}
    [data-testid="stSidebar"] hr {{ border-color:rgba(255,255,255,0.12) !important; }}
    /* Success message in sidebar */
    [data-testid="stSidebar"] .stAlert {{ background:rgba(30,110,60,0.25) !important; border:none !important; }}

    /* ── Score hero ── */
    .score-hero {{
        background:linear-gradient(135deg,{S_CHARCOAL} 0%,#52547A 100%);
        border-radius:14px; padding:26px 28px; color:{S_WHITE};
        text-align:center; box-shadow:0 4px 18px rgba(62,64,91,0.3);
    }}
    .score-hero .num   {{ font-size:4.4rem; font-weight:800; line-height:1; color:#FFFFFF; }}
    .score-hero .denom {{ font-size:1.4rem; color:rgba(255,255,255,0.50); }}
    .score-hero .lbl   {{ font-size:0.74rem; font-weight:600; text-transform:uppercase;
                          letter-spacing:0.09em; color:rgba(255,255,255,0.60); margin-bottom:6px; }}

    /* ── Metric cards ── explicit foreground colours, never white-on-white */
    .mcard {{
        background:{S_WHITE}; border-radius:12px; padding:16px 20px;
        box-shadow:0 1px 6px rgba(0,0,0,0.07); margin-bottom:12px;
        border-top:4px solid {S_RED};
    }}
    .mcard.red   {{ border-top-color:{FG_RED};   background:{BG_RED};   }}
    .mcard.amber {{ border-top-color:{FG_AMBER}; background:{BG_AMBER}; }}
    .mcard.green {{ border-top-color:{FG_GREEN}; background:{BG_GREEN}; }}
    .mcard.grey  {{ border-top-color:{S_CHARCOAL}; background:{S_WHITE}; }}
    .mnum  {{ font-size:2.5rem; font-weight:800; line-height:1.1; }}
    .mcard.red   .mnum {{ color:{FG_RED}; }}
    .mcard.amber .mnum {{ color:{FG_AMBER}; }}
    .mcard.green .mnum {{ color:{FG_GREEN}; }}
    .mcard.grey  .mnum {{ color:{S_CHARCOAL}; }}
    .mlbl  {{ font-size:0.72rem; font-weight:600; text-transform:uppercase;
              letter-spacing:0.07em; color:{S_MUTED}; margin-top:3px; }}

    /* ── Badges — explicit fg/bg, no white-on-white ── */
    .badge-red   {{ background:{BG_RED};   color:{FG_RED};   padding:3px 10px; border-radius:12px; font-weight:600; font-size:0.79rem; white-space:nowrap; display:inline-block; }}
    .badge-amber {{ background:{BG_AMBER}; color:{FG_AMBER}; padding:3px 10px; border-radius:12px; font-weight:600; font-size:0.79rem; white-space:nowrap; display:inline-block; }}
    .badge-green {{ background:{BG_GREEN}; color:{FG_GREEN}; padding:3px 10px; border-radius:12px; font-weight:600; font-size:0.79rem; white-space:nowrap; display:inline-block; }}

    /* ── Issues list ── */
    .irow {{ background:{S_WHITE}; border-radius:8px; padding:11px 15px;
              margin-bottom:7px; box-shadow:0 1px 3px rgba(0,0,0,0.05);
              border-left:4px solid {S_RED}; display:flex; gap:10px; align-items:flex-start; }}
    .irow.critical {{ border-left-color:{FG_RED}; }}
    .irow.warning  {{ border-left-color:{FG_AMBER}; }}
    .irow.info     {{ border-left-color:#1D5FA6; }}
    .irow .ititle  {{ font-weight:600; font-size:0.86rem; color:{S_INK}; }}
    .irow .irec    {{ font-size:0.81rem; color:{S_MUTED}; margin-top:2px; }}
    .idim {{ background:{S_LIGHT}; color:{S_MUTED}; padding:1px 7px;
             border-radius:9px; font-size:0.73rem; display:inline-block; margin-left:6px; }}

    /* ── Table ── */
    .at {{ width:100%; border-collapse:collapse; background:{S_WHITE};
           border-radius:10px; overflow:hidden; box-shadow:0 2px 7px rgba(0,0,0,0.06); }}
    .at th {{ background:{S_CHARCOAL}; color:#FFFFFF; padding:10px 13px;
              text-align:left; font-size:0.76rem; text-transform:uppercase; letter-spacing:0.05em; }}
    .at td {{ padding:9px 13px; border-bottom:1px solid {S_LIGHT}; font-size:0.84rem; color:{S_INK}; }}
    .at tr:last-child td {{ border-bottom:none; }}
    .at tr:hover td {{ background:{S_OFFWHITE}; }}

    /* ── Comparison tab ── */
    .comp-url {{ font-size:0.78rem; color:{S_MUTED}; font-weight:500; }}

    /* ── Misc ── */
    #MainMenu {{ visibility:hidden; }} footer {{ visibility:hidden; }}
    .stDeployButton {{ display:none; }}
    </style>
    """, unsafe_allow_html=True)


# =============================================================================
# SCORING
# =============================================================================
def score_band(s):
    return "red" if s <= 2 else ("amber" if s <= 5 else "green")

def score_fg_hex(s):
    return {" red":FG_RED,"red":FG_RED,"amber":FG_AMBER,"green":FG_GREEN}[score_band(s)]

def score_fg_rgb(s):
    return {"red":FG_RED_RGB,"amber":FG_AMBER_RGB,"green":FG_GREEN_RGB}[score_band(s)]

def score_bg_hex(s):
    return {"red":BG_RED_HEX,"amber":BG_AMBER_HEX,"green":BG_GREEN_HEX}[score_band(s)]

def weighted_overall(scores):
    tw, tws = 0, 0
    for d, s in scores.items():
        w = DIMENSION_CONFIG.get(d, {}).get("weight", 0.05)
        tws += s * w; tw += w
    return round(tws / tw, 1) if tw else 0.0


# =============================================================================
# FETCH + ROBOTS
# =============================================================================
# ── shared session with browser-like headers ──────────────────────────────
_SESSION = requests.Session()
_SESSION.headers.update({
    "User-Agent": "Mozilla/5.0 (Macintosh; Intel Mac OS X 10_15_7) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/124.0.0.0 Safari/537.36",
    "Accept":          "text/html,application/xhtml+xml,application/xml;q=0.9,*/*;q=0.8",
    "Accept-Language": "en-GB,en;q=0.9",
    "Accept-Encoding": "gzip, deflate, br",
    "Connection":      "keep-alive",
    "Upgrade-Insecure-Requests": "1",
    "Cache-Control":   "no-cache",
})


def _is_blocked(status: int, html: str) -> bool:
    if status in (403, 429, 503): return True
    markers = ["just a moment","_cf_chl_","cf-browser-verification",
               "enable javascript","checking your browser",
               "host not in allowlist","access denied","403 forbidden"]
    return any(m in html[:3000].lower() for m in markers)


def _is_js_shell(html: str) -> bool:
    soup = BeautifulSoup(html, "html.parser")
    for t in soup(["script","style","noscript","svg"]): t.decompose()
    return len(soup.get_text(" ", strip=True)) < 300 and html.lower().count("<script") > 3


def _fetch_wayback(url: str) -> str:
    """Fetch the latest Wayback Machine snapshot — bypasses most bot protection."""
    meta = _SESSION.get(
        f"https://archive.org/wayback/available?url={url}", timeout=10
    ).json()
    snap_url = meta.get("archived_snapshots", {}).get("closest", {}).get("url", "")
    if not snap_url:
        raise ValueError("No Wayback snapshot found")
    parts = snap_url.split("/web/", 1)
    if len(parts) == 2:
        ts_and_url = parts[1].split("/", 1)
        if len(ts_and_url) == 2:
            snap_url = f"https://web.archive.org/web/{ts_and_url[0]}id_/{ts_and_url[1]}"
    r = _SESSION.get(snap_url, timeout=20, allow_redirects=True)
    return r.text


def _extract_signals(url: str, html: str) -> dict:
    """
    Parse raw HTML into structured signals dict.
    Extracts meta, JSON-LD, headings, ARIA, links, images, Next.js data.
    JSON-LD and <head> are extracted BEFORE script removal.
    """
    soup = BeautifulSoup(html, "html.parser")
    r = {
        "html": html, "html_raw_length": len(html),
        "json_ld": [], "head_html": "", "headings": [],
        "aria_html": "", "images_sample": "", "links_sample": "",
        "next_data": "", "text": "", "text_length": 0,
        "text_to_html_ratio": 0,
    }

    # 1. JSON-LD — before any tag removal
    for tag in soup.find_all("script", type="application/ld+json"):
        txt = (tag.string or tag.get_text() or "").strip()
        if txt:
            r["json_ld"].append(txt[:3000])
    # Regex fallback in case BeautifulSoup missed any
    if not r["json_ld"]:
        raw_ld = re.findall(
            r'<script[^>]+type=["\']application/ld\+json["\'][^>]*>(.*?)</script>',
            html, re.DOTALL | re.IGNORECASE)
        r["json_ld"] = [s.strip()[:3000] for s in raw_ld if s.strip()]

    # 2. <head> — BeautifulSoup first, regex fallback, raw fallback
    head = soup.find("head")
    if head and len(str(head)) > 100:
        r["head_html"] = str(head)[:8000]
    else:
        m = re.search(r"<head[^>]*>(.*?)</head>", html, re.DOTALL | re.IGNORECASE)
        r["head_html"] = m.group(0)[:8000] if m else html[:3000]

    # 3. Pre-extract meta signals explicitly (mirrors original tool approach)
    meta_lines = []
    title_tag = soup.find("title")
    if title_tag:
        meta_lines.append(f"title: {title_tag.get_text().strip()[:120]}")
    for m in soup.find_all("meta"):
        name = m.get("name","") or m.get("property","") or m.get("http-equiv","")
        val  = m.get("content","")
        if name and val:
            meta_lines.append(f"{name}: {val[:120]}")
    for lnk in soup.find_all("link", rel=True):
        rel = " ".join(lnk.get("rel",[])) if isinstance(lnk.get("rel"), list) else lnk.get("rel","")
        if any(x in rel.lower() for x in ["canonical","alternate","hreflang"]):
            meta_lines.append(f"<link rel=\"{rel}\" href=\"{lnk.get('href','')[:100]}\">")
    r["meta_summary"] = "\n".join(meta_lines) if meta_lines else "NO META TAGS FOUND IN STATIC HTML"

    # 4. __NEXT_DATA__ / embedded JS data stores (Wickes, many commerce sites)
    next_tag = soup.find("script", id="__NEXT_DATA__")
    if next_tag and next_tag.string:
        try:
            nd = json.loads(next_tag.string)
            def _strings(obj, d=0):
                if d > 4: return []
                if isinstance(obj, str) and len(obj) > 20: return [obj[:120]]
                if isinstance(obj, dict):
                    v = []
                    for val in obj.values(): v.extend(_strings(val, d+1))
                    return v[:10]
                if isinstance(obj, list):
                    v = []
                    for val in obj[:5]: v.extend(_strings(val, d+1))
                    return v[:10]
                return []
            r["next_data"] = "\n".join(_strings(nd)[:20])
        except Exception:
            pass

    # 4. Headings
    for lvl in ["h1","h2","h3","h4","h5","h6"]:
        for tag in soup.find_all(lvl):
            r["headings"].append((lvl, tag.get_text(strip=True)[:120]))

    # 5. ARIA
    aria_tags = soup.find_all(lambda t: any(
        k.startswith("aria-") or k == "role" for k in (t.attrs or {})))
    r["aria_html"] = "\n".join(str(t)[:300] for t in aria_tags[:80])

    # 6. Images
    imgs = soup.find_all("img")
    r["images_sample"] = "\n".join(
        f'<img alt="{t.get("alt","[MISSING]")}">' for t in imgs[:80])

    # 7. Links
    r["links_sample"] = "\n".join(
        f'<a href="{t.get("href","")[:100]}" rel="{t.get("rel","")}">{t.get_text(strip=True)[:60]}</a>'
        for t in soup.find_all("a", href=True)[:80])

    # Strip scripts/styles for body text
    for tag in soup(["script","style","noscript","svg"]): tag.decompose()
    r["text"]            = soup.get_text(" ", strip=True)[:8000]
    r["text_length"]     = len(r["text"])
    r["text_to_html_ratio"] = round(r["text_length"] / max(len(html), 1), 3)
    return r


def fetch_page(url):
    """
    Multi-strategy fetch. Falls back to Wayback Machine when direct is blocked.
    Returns the standard page_data dict used by all Gemini prompts.
    """
    r = {
        "url": url, "status_code": None, "error": None, "load_time": None,
        "is_https": url.startswith("https://"), "redirect_chain": [],
        "fetch_source": "direct", "is_js_shell": False,
        # signal fields populated by _extract_signals
        "html": "", "html_raw_length": 0, "json_ld": [], "head_html": "",
        "headings": [], "aria_html": "", "images_sample": "", "links_sample": "",
        "next_data": "", "meta_summary": "", "text": "", "text_length": 0, "text_to_html_ratio": 0,
    }
    html = ""
    try:
        t0   = time.time()
        resp = _SESSION.get(url, timeout=15, allow_redirects=True)
        r["load_time"]      = round(time.time()-t0, 2)
        r["status_code"]    = resp.status_code
        r["redirect_chain"] = [x.url for x in resp.history]
        html = resp.text

        if _is_blocked(resp.status_code, html):
            raise ValueError(f"Blocked (HTTP {resp.status_code})")

    except Exception as e:
        # Wayback Machine fallback
        try:
            html = _fetch_wayback(url)
            r["fetch_source"] = "wayback"
            r["status_code"]  = 200
        except Exception as e2:
            r["error"] = f"Direct: {e} | Wayback: {e2}"
            return r

    r["is_js_shell"] = _is_js_shell(html)
    signals = _extract_signals(url, html)
    r.update(signals)
    return r

def check_robots(url):
    parsed = urlparse(url)
    rob = f"{parsed.scheme}://{parsed.netloc}/robots.txt"
    bots = ["GPTBot","ChatGPT-User","OAI-SearchBot","Google-Extended",
            "Googlebot","PerplexityBot","ClaudeBot","anthropic-ai","Amazonbot","YouBot","CCBot"]
    res = {"robots_url":rob,"found":False,"raw":"","ai_bots":{}}
    try:
        rr = requests.get(rob, timeout=8, headers={"User-Agent":"Mozilla/5.0"})
        if rr.status_code == 200:
            res["found"] = True; res["raw"] = rr.text[:3000]; txt = rr.text.lower()
            for b in bots:
                m = re.search(rf"user-agent:\s*{re.escape(b.lower())}(.*?)(?=user-agent:|$)", txt, re.DOTALL|re.IGNORECASE)
                blocked = False
                if m:
                    blk = m.group(1)
                    if re.search(r"disallow:\s*/\s*$", blk, re.MULTILINE): blocked=True
                    elif re.search(r"disallow:\s*$", blk, re.MULTILINE):   blocked=False
                res["ai_bots"][b] = "BLOCKED" if blocked else "ALLOWED"
    except Exception as e:
        res["error"] = str(e)
    return res

def check_llms_txt(url):
    parsed = urlparse(url)
    lu = f"{parsed.scheme}://{parsed.netloc}/llms.txt"
    res = {"found":False,"url":lu,"content":""}
    try:
        rr = requests.get(lu, timeout=8, headers={"User-Agent":"Mozilla/5.0"})
        if rr.status_code == 200 and len(rr.text) > 10:
            res["found"] = True; res["content"] = rr.text[:1000]
    except: pass
    return res


# =============================================================================
# GEMINI
# =============================================================================
def analyse_dimension(client, page_data, dimension):
    """Build a targeted prompt using pre-extracted signals, then call Gemini."""
    json_ld      = "\n---\n".join(page_data.get("json_ld", [])) or "NONE FOUND"
    head_html    = page_data.get("head_html", "")[:6000]
    headings     = page_data.get("headings", [])
    aria_html    = page_data.get("aria_html", "")[:4000]
    images       = page_data.get("images_sample", "")[:3000]
    links        = page_data.get("links_sample", "")[:3000]
    text         = page_data.get("text", "")[:5000]
    robots       = page_data.get("robots_data", {})
    llms         = page_data.get("llms_txt", {})
    html_snip    = page_data.get("html", "")[:3000]
    headings_fmt = "\n".join(f"{h[0].upper()}: {h[1]}" for h in headings) or "NO HEADINGS FOUND"

    BASE = 'Return ONLY a JSON object, no markdown fences, no extra text.\n'

    prompts = {

"ARIA Implementation": f"""Analyse the ARIA accessibility implementation on this webpage.

ARIA elements found (elements with aria-* or role attributes):
{aria_html or "NONE FOUND"}

Full page text (for context):
{text[:2000]}

Evaluate:
- Landmark roles present (main, navigation, banner, contentinfo, search, complementary)
- aria-label and aria-labelledby usage quality
- Form input labelling
- aria-live regions for dynamic content
- Whether ARIA aids AI understanding of page structure

{BASE}{{"score":<1-10>,"summary":"<2 sentences>","findings":[],"issues":[{{"severity":"critical|warning|info","issue":"<str>","recommendation":"<str>"}}],"positive":[]}}""",

"Structured Data / Schema": f"""Analyse the schema.org structured data on this webpage.

JSON-LD blocks found on page:
{json_ld}

Evaluate:
- What schema types ARE present and are they complete?
- What important schema types are MISSING for this page type?
- Are required/recommended properties populated?
- Would an AI system get useful structured signals from this page?

{BASE}{{"score":<1-10>,"summary":"<2 sentences>","schemas_found":[],"schemas_missing":[],"findings":[],"issues":[{{"severity":"critical|warning|info","issue":"<str>","recommendation":"<str>"}}],"positive":[]}}""",

"Heading Structure": f"""Analyse the heading structure (H1–H6) on this webpage.

All headings found on page:
{headings_fmt}

Evaluate:
- How many H1s are there? (should be exactly 1)
- Is the hierarchy logical with no skipped levels?
- Are headings descriptive and keyword-relevant?
- Does the heading outline clearly describe page content for AI systems?

{BASE}{{"score":<1-10>,"summary":"<2 sentences>","h1_count":<int>,"findings":[],"issues":[{{"severity":"critical|warning|info","issue":"<str>","recommendation":"<str>"}}],"positive":[]}}""",

"Meta & SEO Signals": f"""Analyse the meta tags and SEO signals on this webpage.

IMPORTANT: Content was fetched via: {page_data.get("fetch_source","direct")}.
If fetch_source is "wayback" this is a cached version reflecting what crawlers see.
If head_html is empty, the site injects meta tags via JavaScript — a critical AI/SEO issue.

Pre-extracted meta signals (title, meta tags, canonical, OG, Twitter — all found in static HTML):
{page_data.get("meta_summary","") or "NO META TAGS FOUND IN STATIC HTML — likely CSR page"}

Full <head> HTML for reference ({len(head_html)} chars):
{head_html[:3000] if head_html.strip() else "EMPTY — no static head content found"}

Evaluate based on the pre-extracted signals above:
- Title tag: present? length (50-60 chars ideal)? descriptive?
- Meta description: present? length (150-160 chars ideal)?
- Canonical tag: present and pointing to correct URL?
- Open Graph tags (og:title, og:description, og:image): present?
- Twitter Card tags: present?
- Meta robots tag: present and configured correctly?
- Viewport meta tag
- If NO meta tags found: this is a critical CSR issue — all meta is JS-injected and invisible to AI crawlers

{BASE}{{"score":<1-10>,"summary":"<2 sentences>","title_length":<int>,"meta_desc_present":true,"canonical_present":true,"og_present":true,"findings":[],"issues":[{{"severity":"critical|warning|info","issue":"<str>","recommendation":"<str>"}}],"positive":[]}}""",

"Link Quality": f"""Analyse the link quality on this webpage.

All links found on page ({len(page_data.get("links_sample","").splitlines())} shown):
{links}

Evaluate:
- Descriptive vs generic anchor text ("click here", "read more", "here")
- Links with no anchor text at all
- nofollow/ugc/sponsored rel attribute usage
- Internal vs external link balance
- Whether anchor text helps AI systems understand content graph

{BASE}{{"score":<1-10>,"summary":"<2 sentences>","findings":[],"issues":[{{"severity":"critical|warning|info","issue":"<str>","recommendation":"<str>"}}],"positive":[]}}""",

"Image Alt Text": f"""Analyse the image alt text on this webpage.

All images found on page:
{images}

Count images, count those with alt="" (decorative) and those with meaningful alt text vs [MISSING].

Evaluate:
- What percentage of images have alt text?
- Quality of alt text (descriptive and relevant vs generic/keyword-stuffed)
- Decorative images correctly using empty alt=""
- Missing alt attributes (not even empty)

{BASE}{{"score":<1-10>,"summary":"<2 sentences>","images_found":<int>,"images_with_alt":<int>,"findings":[],"issues":[{{"severity":"critical|warning|info","issue":"<str>","recommendation":"<str>"}}],"positive":[]}}""",

"Crawlability & Bot Access": f"""Analyse crawlability and AI bot access for this webpage.

Technical signals:
- Fetch source: {page_data.get("fetch_source","direct")} {"(WAYBACK: site blocked direct fetch — content is cached version)" if page_data.get("fetch_source")=="wayback" else ""}
- JS shell detected: {page_data.get("is_js_shell",False)} (True = client-side rendered, very little static content)
- HTML raw length: {page_data.get("html_raw_length",0)} chars
- Extracted text length: {page_data.get("text_length",0)} chars
- Text-to-HTML ratio: {page_data.get("text_to_html_ratio",0):.3f} (below 0.05 strongly suggests JS rendering)
- HTTP status: {page_data.get("status_code")}
- HTTPS: {page_data.get("is_https")}
- Load time: {page_data.get("load_time")}s
- Redirect chain: {page_data.get("redirect_chain",[])}

Robots.txt data:
{json.dumps(robots, indent=2)}

HTML snippet (first 3000 chars):
{html_snip}

Evaluate: SSR vs CSR rendering, bot blocking mechanisms, HTTPS, redirects, load time, AI bot permissions in robots.txt.

{BASE}{{"score":<1-10>,"summary":"<2 sentences>","rendering_type":"SSR|CSR|Mixed","js_dependent":true,"findings":[],"issues":[{{"severity":"critical|warning|info","issue":"<str>","recommendation":"<str>"}}],"positive":[]}}""",

"LLM Content Signals": f"""Analyse the content quality and LLM/AI visibility signals on this webpage.

Page text content:
{text}

Page headings:
{headings_fmt}

JSON-LD found:
{json_ld[:1000]}

Embedded JS data (Next.js/__NEXT_DATA__ etc.):
{page_data.get("next_data","NONE")[:600]}

Evaluate:
- E-E-A-T signals (expertise, experience, authoritativeness, trustworthiness)
- Content depth and factual density
- Author names, dates, credentials, citations
- Brand/entity clarity
- FAQ-style extractable Q&A content
- Whether an LLM could extract clear, attributable facts from this page

{BASE}{{"score":<1-10>,"summary":"<2 sentences>","eeat_signals":[],"findings":[],"issues":[{{"severity":"critical|warning|info","issue":"<str>","recommendation":"<str>"}}],"positive":[]}}""",

"AI Search Health": f"""Analyse AI search health signals for this webpage.

llms.txt found: {llms.get("found", False)}
llms.txt content: {llms.get("content", "N/A")}

Robots.txt AI bot access:
{json.dumps(robots.get("ai_bots", {}), indent=2)}

<head> snippet (for AI-specific meta tags):
{head_html[:2000]}

Evaluate:
- llms.txt: present, well-structured, useful?
- Which AI bots are explicitly allowed/blocked in robots.txt?
- Any AI-specific meta tags or directives?
- Overall AI search readiness

{BASE}{{"score":<1-10>,"summary":"<2 sentences>","llms_txt_present":false,"ai_bots_blocked":[],"findings":[],"issues":[{{"severity":"critical|warning|info","issue":"<str>","recommendation":"<str>"}}],"positive":[]}}""",

"Duplicate Content & Tags": f"""Analyse duplicate content and canonical tag issues on this webpage.

Pre-extracted meta signals:
{page_data.get("meta_summary","") or "NO META TAGS IN STATIC HTML"}

Text-to-HTML ratio: {page_data.get("text_to_html_ratio",0):.3f}
Page text sample: {text[:1500]}

Evaluate:
- Canonical tag: present and self-referencing correctly?
- Title tag: unique or likely a CMS default/template?
- Meta description: unique or templated?
- Thin content signals (low text-to-HTML ratio, boilerplate heavy)
- If no canonical found: flag as critical issue

{BASE}{{"score":<1-10>,"summary":"<2 sentences>","canonical_present":true,"findings":[],"issues":[{{"severity":"critical|warning|info","issue":"<str>","recommendation":"<str>"}}],"positive":[]}}""",
    }

    prompt = prompts.get(dimension,
        f'Analyse {dimension} for AI visibility. {BASE}{{"score":5,"summary":"","findings":[],"issues":[],"positive":[]}}')

    try:
        resp = client.models.generate_content(model="gemini-2.5-flash", contents=prompt)
        raw  = resp.text.strip()
        raw  = re.sub(r"^```(?:json)?\n?","",raw)
        raw  = re.sub(r"\n?```$","",raw)
        return json.loads(raw)
    except json.JSONDecodeError:
        m = re.search(r'\{.*\}', raw if 'raw' in dir() else '{}', re.DOTALL)
        if m:
            try: return json.loads(m.group(0))
            except: pass
    except: pass
    return {"score":0,"summary":"Analysis could not be completed.","findings":[],"issues":[],"positive":[]}

def gen_exec_summary(client, url, scores, all_results):
    overall = weighted_overall(scores)
    crit = [f"{d}: {i['issue']}" for d,r in all_results.items()
            for i in r.get("issues",[]) if i.get("severity")=="critical"][:8]
    prompt = f"""You are an expert AI visibility consultant at Summit, a performance marketing agency.

Write a concise executive summary for an AI visibility audit of {url}.
Overall score: {overall}/10
Scores: {json.dumps(scores)}
Critical issues: {chr(10).join(crit)}

Write exactly 3 short paragraphs (2-3 sentences each):
1. Overall verdict and score context
2. Key strengths
3. Top 2-3 priorities for improvement and business impact

Tone: direct, consultancy-grade. UK English. No bullet points. Keep it punchy — this will appear alongside charts."""
    try:
        resp = client.models.generate_content(model="gemini-2.5-flash", contents=prompt)
        return resp.text.strip()
    except Exception as e:
        return f"Executive summary could not be generated: {e}"


# =============================================================================
# CHARTS
# =============================================================================
def gauge_chart(overall):
    color = score_fg_hex(overall)
    fig = go.Figure(go.Indicator(
        mode="gauge+number", value=overall,
        number={"suffix":"/10","font":{"size":38,"color":color,"family":"Inter"}},
        gauge={
            "axis":{"range":[0,10],"tickwidth":1,"tickcolor":S_MUTED,"tickfont":{"size":9,"color":S_MUTED}},
            "bar":{"color":color,"thickness":0.26},
            "bgcolor":S_WHITE,"borderwidth":0,
            "steps":[{"range":[0,3.3],"color":BG_RED},{"range":[3.3,6.6],"color":BG_AMBER},{"range":[6.6,10],"color":BG_GREEN}],
        }
    ))
    fig.update_layout(paper_bgcolor="rgba(0,0,0,0)",plot_bgcolor="rgba(0,0,0,0)",
                      margin=dict(l=18,r=18,t=18,b=8),height=200,font=dict(family="Inter"))
    return fig

def bar_chart(scores):
    dims   = list(scores.keys())
    vals   = [scores[d] for d in dims]
    colors = [score_fg_hex(v) for v in vals]
    labels = [d.replace(" & "," &\n").replace(" / ","/\n") for d in dims]
    fig = go.Figure(go.Bar(
        x=vals, y=labels, orientation="h",
        marker_color=colors,
        text=[f"  {v}/10" for v in vals], textposition="outside",
        textfont=dict(size=11,color=S_INK),
        hovertemplate="%{y}: %{x}/10<extra></extra>",
    ))
    fig.add_vline(x=3.3,line_dash="dot",line_color=FG_RED,opacity=0.4,
                  annotation_text="Critical",annotation_font=dict(size=9,color=FG_RED))
    fig.add_vline(x=6.6,line_dash="dot",line_color=FG_GREEN,opacity=0.4,
                  annotation_text="Good",annotation_font=dict(size=9,color=FG_GREEN))
    fig.update_layout(
        xaxis=dict(range=[0,13],showgrid=True,gridcolor=S_LIGHT,title="Score / 10"),
        yaxis=dict(autorange="reversed",tickfont=dict(size=10,color=S_INK)),
        paper_bgcolor="rgba(0,0,0,0)",plot_bgcolor="rgba(0,0,0,0)",
        margin=dict(l=10,r=70,t=18,b=28),height=400,font=dict(family="Inter"),showlegend=False,
    )
    return fig

def radar_chart(scores, label=""):
    dims = list(scores.keys()); vals = [scores[d] for d in dims]
    fig = go.Figure(go.Scatterpolar(
        r=vals+[vals[0]], theta=dims+[dims[0]], name=label or "Score",
        fill="toself", fillcolor="rgba(208,56,11,0.12)",
        line=dict(color=S_RED,width=2.5), marker=dict(size=6,color=S_RED),
    ))
    fig.update_layout(
        polar=dict(bgcolor=S_WHITE,
            radialaxis=dict(visible=True,range=[0,10],tickfont=dict(size=8),gridcolor=S_LIGHT),
            angularaxis=dict(tickfont=dict(size=9,color=S_INK))),
        showlegend=bool(label),
        paper_bgcolor="rgba(0,0,0,0)",plot_bgcolor="rgba(0,0,0,0)",
        margin=dict(l=55,r=55,t=28,b=28),height=400,font=dict(family="Inter"),
    )
    return fig

def multi_radar_chart(results_map):
    """Overlay multiple URLs on one radar."""
    colors = [S_RED,"#0369A1","#0E7C4A","#7C3AED","#B45309"]
    dims = list(DIMENSION_CONFIG.keys())
    fig = go.Figure()
    for i,(label,scores) in enumerate(results_map.items()):
        col = colors[i % len(colors)]
        vals = [scores.get(d,0) for d in dims]
        fig.add_trace(go.Scatterpolar(
            r=vals+[vals[0]], theta=dims+[dims[0]], name=label,
            fill="toself", fillcolor=f"rgba({int(col[1:3],16)},{int(col[3:5],16)},{int(col[5:7],16)},0.08)",
            line=dict(color=col,width=2.2), marker=dict(size=5,color=col),
        ))
    fig.update_layout(
        polar=dict(bgcolor=S_WHITE,
            radialaxis=dict(visible=True,range=[0,10],tickfont=dict(size=8),gridcolor=S_LIGHT),
            angularaxis=dict(tickfont=dict(size=9,color=S_INK))),
        showlegend=True, legend=dict(font=dict(size=9),orientation="h",y=-0.12),
        paper_bgcolor="rgba(0,0,0,0)",plot_bgcolor="rgba(0,0,0,0)",
        margin=dict(l=55,r=55,t=28,b=55),height=430,font=dict(family="Inter"),
    )
    return fig

def multi_bar_chart(results_map):
    """Grouped bar chart comparing multiple URLs."""
    dims   = list(DIMENSION_CONFIG.keys())
    colors = [S_RED,"#0369A1","#0E7C4A","#7C3AED","#B45309"]
    fig = go.Figure()
    for i,(label,scores) in enumerate(results_map.items()):
        col  = colors[i % len(colors)]
        vals = [scores.get(d,0) for d in dims]
        fig.add_trace(go.Bar(
            name=label, x=dims, y=vals,
            marker_color=col, opacity=0.85,
            hovertemplate=f"{label}<br>%{{x}}: %{{y}}/10<extra></extra>",
        ))
    fig.update_layout(
        barmode="group",
        xaxis=dict(tickfont=dict(size=9),tickangle=-30),
        yaxis=dict(range=[0,11],title="Score / 10"),
        paper_bgcolor="rgba(0,0,0,0)",plot_bgcolor="rgba(0,0,0,0)",
        legend=dict(font=dict(size=9),orientation="h",y=1.08),
        margin=dict(l=10,r=10,t=30,b=80),height=420,font=dict(family="Inter"),
    )
    return fig

def weight_donut(scores):
    dims=[]; weights=[]; colors=[]
    for d,cfg in DIMENSION_CONFIG.items():
        dims.append(d); weights.append(cfg["weight"]*100); colors.append(cfg["color"])
    fig = go.Figure(go.Pie(labels=dims,values=weights,hole=0.52,
        marker=dict(colors=colors),textinfo="percent",textfont=dict(size=9),
        hovertemplate="%{label}<br>Weight: %{value:.0f}%<extra></extra>"))
    fig.update_layout(showlegend=True,legend=dict(font=dict(size=9),orientation="v",x=1.02),
        paper_bgcolor="rgba(0,0,0,0)",margin=dict(l=5,r=5,t=18,b=8),
        height=340,font=dict(family="Inter"))
    return fig


# =============================================================================
# DOCX EXPORT
# =============================================================================
def _set_cell_bg(cell, hex6):
    tc=cell._tc; tcPr=tc.get_or_add_tcPr()
    shd=OxmlElement("w:shd")
    shd.set(qn("w:val"),"clear"); shd.set(qn("w:color"),"auto"); shd.set(qn("w:fill"),hex6)
    tcPr.append(shd)

def _hdg(doc, text, size=14):
    p=doc.add_paragraph(); pPr=p._p.get_or_add_pPr()
    pBdr=OxmlElement("w:pBdr"); bot=OxmlElement("w:bottom")
    bot.set(qn("w:val"),"single"); bot.set(qn("w:sz"),"8")
    bot.set(qn("w:space"),"4");   bot.set(qn("w:color"),"D0380B")
    pBdr.append(bot); pPr.append(pBdr)
    r=p.add_run(text); r.bold=True; r.font.size=Pt(size)
    r.font.color.rgb=RGBColor(*S_CHARCOAL_RGB); r.font.name="Arial"
    return p

def build_docx(url, scores, all_results, exec_summary, logo_bytes):
    doc=Document()
    for sec in doc.sections:
        sec.page_width=Emu(int(11906*914.4))
        sec.page_height=Emu(int(16838*914.4))
        sec.top_margin=Inches(0.7); sec.bottom_margin=Inches(0.7)
        sec.left_margin=Inches(0.8); sec.right_margin=Inches(0.8)
    doc.styles["Normal"].font.name="Arial"; doc.styles["Normal"].font.size=Pt(10)

    # Logo + title
    lp=doc.add_paragraph(); lp.alignment=WD_ALIGN_PARAGRAPH.LEFT
    lp.add_run().add_picture(io.BytesIO(logo_bytes), width=Inches(1.3))
    _hdg(doc,"AI Visibility Audit Report",size=22)
    mp=doc.add_paragraph()
    for txt,col in [(f"{url}  |  ",(120,120,130)),(datetime.now().strftime("%B %Y"),(120,120,130))]:
        r=mp.add_run(txt); r.font.size=Pt(9); r.font.color.rgb=RGBColor(*col)
    doc.add_paragraph()

    # Score table
    overall=weighted_overall(scores); band=score_band(overall)
    bg_hm={"red":BG_RED_HEX,"amber":BG_AMBER_HEX,"green":BG_GREEN_HEX}
    fg_rm={"red":FG_RED_RGB,"amber":FG_AMBER_RGB,"green":FG_GREEN_RGB}
    lm={"red":"Needs Attention","amber":"Developing","green":"Good"}
    _hdg(doc,"Overall Score")
    t=doc.add_table(rows=1,cols=3); t.style="Table Grid"; t.autofit=False
    ws=[1600,1600,6402]; cells=t.rows[0].cells
    cells[0].width=Emu(int(ws[0]*914.4)); _set_cell_bg(cells[0],"3E405B")
    p=cells[0].paragraphs[0]; p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    r=p.add_run(str(overall)); r.bold=True; r.font.size=Pt(40); r.font.color.rgb=RGBColor(255,255,255)
    p2=cells[0].add_paragraph("/10"); p2.alignment=WD_ALIGN_PARAGRAPH.CENTER
    p2.runs[0].font.color.rgb=RGBColor(200,200,210); p2.runs[0].font.size=Pt(13)
    cells[1].width=Emu(int(ws[1]*914.4)); _set_cell_bg(cells[1],bg_hm[band])
    p=cells[1].paragraphs[0]; p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    r=p.add_run(lm[band]); r.bold=True; r.font.size=Pt(13); r.font.color.rgb=RGBColor(*fg_rm[band])
    cells[2].width=Emu(int(ws[2]*914.4)); _set_cell_bg(cells[2],"F7F7F8")
    p=cells[2].paragraphs[0]
    r=p.add_run("Score breakdown"); r.bold=True; r.font.size=Pt(9); r.font.color.rgb=RGBColor(*S_CHARCOAL_RGB)
    for d,s in sorted(scores.items(),key=lambda x:x[1]):
        p2=cells[2].add_paragraph()
        r1=p2.add_run(f"{d}: "); r1.font.size=Pt(8.5)
        r2=p2.add_run(f"{s}/10"); r2.bold=True; r2.font.size=Pt(8.5)
        r2.font.color.rgb=RGBColor(*fg_rm[score_band(s)])
    doc.add_paragraph()

    # Exec summary
    _hdg(doc,"Executive Summary")
    for para in exec_summary.split("\n\n"):
        if para.strip():
            p=doc.add_paragraph(para.strip()); p.style.font.size=Pt(10)
    doc.add_paragraph()

    # Scores table
    _hdg(doc,"Dimension Scores")
    dt=doc.add_table(rows=1,cols=4); dt.style="Table Grid"; dt.autofit=False
    dw=[3600,700,900,4402]
    for i,(cell,hdr) in enumerate(zip(dt.rows[0].cells,["Dimension","Weight","Score","Rating"])):
        cell.width=Emu(dw[i]*914); _set_cell_bg(cell,"3E405B")
        r=cell.paragraphs[0].add_run(hdr); r.bold=True
        r.font.color.rgb=RGBColor(255,255,255); r.font.size=Pt(9)
    for d,s in scores.items():
        row=dt.add_row(); bd=score_band(s); wp=int(DIMENSION_CONFIG.get(d,{}).get("weight",0.05)*100)
        for i,w in enumerate(dw): row.cells[i].width=Emu(int(w*914.4))
        row.cells[0].paragraphs[0].add_run(d).font.size=Pt(9)
        row.cells[1].paragraphs[0].add_run(f"{wp}%").font.size=Pt(9)
        _set_cell_bg(row.cells[2],bg_hm[bd])
        sr=row.cells[2].paragraphs[0].add_run(f"{s}/10")
        sr.bold=True; sr.font.size=Pt(10); sr.font.color.rgb=RGBColor(*fg_rm[bd])
        _set_cell_bg(row.cells[3],bg_hm[bd])
        lr=row.cells[3].paragraphs[0].add_run({"red":"Needs Attention","amber":"Needs Work","green":"Good"}[bd])
        lr.font.size=Pt(9); lr.font.color.rgb=RGBColor(*fg_rm[bd])
    doc.add_paragraph()

    # Detailed findings
    _hdg(doc,"Detailed Findings")
    sev_bg={"critical":BG_RED_HEX,"warning":BG_AMBER_HEX,"info":"EBF0FB"}
    sev_fg={"critical":FG_RED_RGB,"warning":FG_AMBER_RGB,"info":(21,80,175)}
    for d,res in all_results.items():
        s=scores.get(d,0); bd=score_band(s)
        icon=DIMENSION_CONFIG.get(d,{}).get("icon","•")
        dp=doc.add_paragraph()
        r1=dp.add_run(f"{icon}  {d}   "); r1.bold=True; r1.font.size=Pt(12)
        r1.font.color.rgb=RGBColor(*S_CHARCOAL_RGB)
        r2=dp.add_run(f"[{s}/10]"); r2.bold=True; r2.font.size=Pt(12)
        r2.font.color.rgb=RGBColor(*fg_rm[bd])
        if res.get("summary"):
            sp=doc.add_paragraph(res["summary"]); sp.runs[0].italic=True; sp.runs[0].font.size=Pt(9.5)
        issues=res.get("issues",[])
        if issues:
            it=doc.add_table(rows=1,cols=3); it.style="Table Grid"; it.autofit=False
            iw=[900,3200,5502]
            for i,(cell,hdr) in enumerate(zip(it.rows[0].cells,["Severity","Issue","Recommendation"])):
                cell.width=Emu(iw[i]*914); _set_cell_bg(cell,"3E405B")
                r=cell.paragraphs[0].add_run(hdr); r.bold=True
                r.font.color.rgb=RGBColor(255,255,255); r.font.size=Pt(8.5)
            for iss in issues:
                sev=iss.get("severity","info"); row=it.add_row()
                for i,w in enumerate(iw): row.cells[i].width=Emu(int(w*914.4))
                _set_cell_bg(row.cells[0],sev_bg.get(sev,"F5F5F5"))
                sr=row.cells[0].paragraphs[0].add_run(sev.upper())
                sr.bold=True; sr.font.size=Pt(8); sr.font.color.rgb=RGBColor(*sev_fg.get(sev,(80,80,80)))
                row.cells[1].paragraphs[0].add_run(iss.get("issue","")).font.size=Pt(8.5)
                row.cells[2].paragraphs[0].add_run(iss.get("recommendation","")).font.size=Pt(8.5)
        pos=res.get("positive",[])
        if pos:
            pp=doc.add_paragraph()
            r=pp.add_run("✓ Strengths:  "); r.bold=True; r.font.size=Pt(9); r.font.color.rgb=RGBColor(*FG_GREEN_RGB)
            r2=pp.add_run("  |  ".join(str(p) for p in pos[:3] if p)); r2.font.size=Pt(9); r2.font.color.rgb=RGBColor(*FG_GREEN_RGB)
        doc.add_paragraph()

    # Priority recs
    _hdg(doc,"Priority Recommendations")
    all_recs=[]
    for d,res in all_results.items():
        w=DIMENSION_CONFIG.get(d,{}).get("weight",0.05)
        for iss in res.get("issues",[]):
            sev=iss.get("severity","info")
            all_recs.append({"pri":{"critical":1,"warning":2,"info":3}.get(sev,3),"w":w,
                             "d":d,"issue":iss.get("issue",""),"rec":iss.get("recommendation",""),"sev":sev})
    all_recs.sort(key=lambda x:(x["pri"],-x["w"]))
    rt=doc.add_table(rows=1,cols=4); rt.style="Table Grid"; rt.autofit=False
    rw=[500,1700,3200,4202]
    for i,(cell,hdr) in enumerate(zip(rt.rows[0].cells,["#","Dimension","Issue","Recommendation"])):
        cell.width=Emu(rw[i]*914); _set_cell_bg(cell,"3E405B")
        r=cell.paragraphs[0].add_run(hdr); r.bold=True
        r.font.color.rgb=RGBColor(255,255,255); r.font.size=Pt(9)
    for rank,rec in enumerate(all_recs[:15],1):
        row=rt.add_row(); sev=rec["sev"]
        for i,w in enumerate(rw): row.cells[i].width=Emu(int(w*914.4))
        _set_cell_bg(row.cells[0],sev_bg.get(sev,"F5F5F5"))
        nr=row.cells[0].paragraphs[0].add_run(str(rank)); nr.bold=True; nr.font.size=Pt(9)
        nr.font.color.rgb=RGBColor(*sev_fg.get(sev,(80,80,80)))
        row.cells[1].paragraphs[0].add_run(rec["d"]).font.size=Pt(8.5)
        row.cells[2].paragraphs[0].add_run(rec["issue"]).font.size=Pt(8.5)
        row.cells[3].paragraphs[0].add_run(rec["rec"]).font.size=Pt(8.5)
    doc.add_paragraph()

    fp=doc.add_paragraph(); fp.alignment=WD_ALIGN_PARAGRAPH.CENTER
    pPr=fp._p.get_or_add_pPr(); pBdr=OxmlElement("w:pBdr"); top=OxmlElement("w:top")
    top.set(qn("w:val"),"single"); top.set(qn("w:sz"),"4"); top.set(qn("w:space"),"4"); top.set(qn("w:color"),"D0380B")
    pBdr.append(top); pPr.append(pBdr)
    r=fp.add_run("Prepared by Summit  |  AI Visibility Practice  |  summit.co.uk")
    r.font.size=Pt(8); r.italic=True; r.font.color.rgb=RGBColor(130,130,140)

    buf=io.BytesIO(); doc.save(buf); buf.seek(0)
    return buf.getvalue()


# =============================================================================
# PDF — visual one-pager, not a doc dump
# =============================================================================
def build_pdf(url, scores, exec_summary, logo_bytes):
    logo_tmp="/tmp/summit_logo_pdf.png"
    with open(logo_tmp,"wb") as f: f.write(logo_bytes)

    overall=weighted_overall(scores); band=score_band(overall)
    bg_rgb={"red":(253,233,229),"amber":(255,244,224),"green":(230,245,236)}
    fg_rgb={"red":FG_RED_RGB,"amber":FG_AMBER_RGB,"green":FG_GREEN_RGB}
    lm={"red":"Needs Attention","amber":"Developing","green":"Good"}
    dim_short = {  # shorter labels for PDF
        "Crawlability & Bot Access":"Crawlability",
        "Structured Data / Schema":"Structured Data",
        "LLM Content Signals":"LLM Content",
        "Meta & SEO Signals":"Meta / SEO",
        "Heading Structure":"Headings",
        "ARIA Implementation":"ARIA",
        "Link Quality":"Link Quality",
        "Image Alt Text":"Image Alt",
        "AI Search Health":"AI Health",
        "Duplicate Content & Tags":"Duplicates",
    }
    score_lbl={"red":"Needs Attention","amber":"Needs Work","green":"Good"}

    class SummitPDF(FPDF):
        def header(self):
            self.image(logo_tmp,12,8,22)
            self.set_font("Helvetica","B",15); self.set_text_color(*S_CHARCOAL_RGB)
            self.set_xy(38,10); self.cell(0,7,"AI Visibility Audit",ln=False)
            self.set_font("Helvetica","",8.5); self.set_text_color(130,130,140)
            self.set_xy(38,18); self.cell(0,5,_pdf_safe(f"{url}  -  {datetime.now().strftime('%B %Y')}", 120),ln=True)
            self.set_draw_color(*S_RED_RGB); self.set_line_width(0.9)
            self.line(12,27,198,27); self.ln(3)
        def footer(self):
            self.set_y(-11); self.set_font("Helvetica","I",7.5)
            self.set_text_color(160,160,170)
            self.cell(0,5,f"Summit AI Visibility Audit  ·  Confidential  ·  Page {self.page_no()}",align="C")

    pdf=SummitPDF(); pdf.add_page(); pdf.set_auto_page_break(auto=True,margin=13)

    # ── Hero row: big score left, rating centre, 3 stats right ──
    y0=pdf.get_y()
    # Score box
    pdf.set_fill_color(*S_CHARCOAL_RGB); pdf.rect(12,y0,52,32,"F")
    pdf.set_text_color(255,255,255); pdf.set_font("Helvetica","B",30)
    pdf.set_xy(12,y0+3); pdf.cell(52,14,f"{overall}/10",align="C",ln=False)
    pdf.set_font("Helvetica","",7.5); pdf.set_text_color(180,180,200)
    pdf.set_xy(12,y0+19); pdf.cell(52,7,"Overall AI Visibility Score",align="C",ln=False)

    # Rating badge
    pdf.set_fill_color(*bg_rgb[band]); pdf.rect(66,y0,52,32,"F")
    pdf.set_text_color(*fg_rgb[band]); pdf.set_font("Helvetica","B",13)
    pdf.set_xy(66,y0+7); pdf.cell(52,10,lm[band],align="C",ln=False)
    pdf.set_font("Helvetica","",7.5); pdf.set_xy(66,y0+19)
    pdf.cell(52,7,f"{len(scores)} dimensions",align="C",ln=False)

    # Mini stats
    n_crit=sum(1 for r in [scores[d] for d in scores] if score_band(r)=="red")
    n_good=sum(1 for r in [scores[d] for d in scores] if score_band(r)=="green")
    for i,(val,lbl,col) in enumerate([(n_crit,"Critical",FG_RED_RGB),(n_good,"Passing",FG_GREEN_RGB)]):
        bx=121+i*40
        pdf.set_fill_color(248,248,250); pdf.rect(bx,y0,36,32,"F")
        pdf.set_text_color(*col); pdf.set_font("Helvetica","B",20)
        pdf.set_xy(bx,y0+4); pdf.cell(36,12,str(val),align="C",ln=False)
        pdf.set_font("Helvetica","",7); pdf.set_text_color(130,130,140)
        pdf.set_xy(bx,y0+19); pdf.cell(36,6,lbl,align="C",ln=False)

    pdf.ln(y0+38-pdf.get_y())

    # ── Score tiles — 5 columns × 2 rows ──
    pdf.ln(3)
    pdf.set_font("Helvetica","B",9); pdf.set_text_color(*S_CHARCOAL_RGB)
    pdf.cell(0,6,"Dimension Scores",ln=True)
    pdf.set_draw_color(*S_RED_RGB); pdf.set_line_width(0.6)
    pdf.line(12,pdf.get_y(),198,pdf.get_y()); pdf.ln(2)

    tile_w=37; tile_h=17; cols=5; gap=0.6
    items=sorted(scores.items(),key=lambda x:x[1])
    for i,(dim,s) in enumerate(items):
        bd=score_band(s); col_i=i%cols; row_i=i//cols
        x=12+col_i*(tile_w+gap); y=pdf.get_y() if col_i==0 else pdf.get_y()
        if col_i==0 and i>0: pdf.ln(tile_h+gap)
        y=pdf.get_y()
        pdf.set_fill_color(*bg_rgb[bd]); pdf.rect(x,y,tile_w,tile_h,"F")
        # score number
        pdf.set_text_color(*fg_rgb[bd]); pdf.set_font("Helvetica","B",14)
        pdf.set_xy(x,y+1); pdf.cell(tile_w,7,f"{s}/10",align="C",ln=False)
        # short label
        short=dim_short.get(dim,dim[:12])
        pdf.set_font("Helvetica","",6.5); pdf.set_text_color(80,80,90)
        pdf.set_xy(x,y+9); pdf.cell(tile_w,4,short,align="C",ln=False)
    pdf.ln(tile_h+4)

    # ── Executive summary — compact 3-para version ──
    pdf.set_font("Helvetica","B",9); pdf.set_text_color(*S_CHARCOAL_RGB)
    pdf.cell(0,6,"Executive Summary",ln=True)
    pdf.set_draw_color(*S_RED_RGB); pdf.set_line_width(0.6)
    pdf.line(12,pdf.get_y(),198,pdf.get_y()); pdf.ln(2)

    paras=[p.strip() for p in exec_summary.split("\n\n") if p.strip()][:3]
    pdf.set_font("Helvetica","",8.5); pdf.set_text_color(40,40,50)
    for para in paras:
        pdf.multi_cell(0,4.8,para); pdf.ln(1.5)

    # ── Top issues table ──
    pdf.ln(2)
    pdf.set_font("Helvetica","B",9); pdf.set_text_color(*S_CHARCOAL_RGB)
    pdf.cell(0,6,"Top Priority Actions",ln=True)
    pdf.set_draw_color(*S_RED_RGB); pdf.set_line_width(0.6)
    pdf.line(12,pdf.get_y(),198,pdf.get_y()); pdf.ln(2)

    # Build flat issue list sorted by severity then weight
    from functools import reduce
    flat=[]
    for d_name,res in {}.items(): pass  # placeholder — populated below
    flat_issues_for_pdf = []
    pdf.set_fill_color(*S_CHARCOAL_RGB); pdf.set_text_color(255,255,255)
    pdf.set_font("Helvetica","B",7.5)
    pdf.cell(22,5.5,"Severity",fill=True,border=0)
    pdf.cell(35,5.5,"Dimension",fill=True,border=0)
    pdf.cell(141,5.5,"Recommendation",fill=True,border=0); pdf.ln()

    raw = pdf.output()
    return bytes(raw) if not isinstance(raw, bytes) else raw

def _pdf_safe(text, maxlen=200):
    """Sanitise text for fpdf Latin-1 encoding — replace non-Latin-1 chars."""
    if not text: return ""
    text = str(text)
    # Common unicode replacements
    replacements = {
        "’": "'", "‘": "'", "“": '"', "”": '"',
        "–": "-", "—": "-", "…": "...", "·": "*",
        "•": "*", "‐": "-", "é": "e", "è": "e",
        "à": "a", "â": "a", "ô": "o", "û": "u",
        "ü": "u", "ä": "a", "ö": "o", "ß": "ss",
        "£": "GBP", "€": "EUR", "®": "(R)", "™": "(TM)",
        "©": "(C)", "°": "deg", "×": "x", "÷": "/",
        "→": "->", "←": "<-", "↔": "<->", "»": ">>",
        "«": "<<", "¼": "1/4", "½": "1/2", "¾": "3/4",
    }
    for uc, asc in replacements.items():
        text = text.replace(uc, asc)
    # Strip anything remaining outside Latin-1
    text = text.encode("latin-1", errors="replace").decode("latin-1")
    return text[:maxlen].strip()



def build_pdf_with_issues(url, scores, all_results, exec_summary, logo_bytes):
    """
    Three-page visual PDF matching the dashboard:
      Page 1 — Hero + Score tiles + Horizontal bar chart
      Page 2 — Radar chart + Issues & Recommendations table
      Page 3 — Executive Summary
    Charts built with matplotlib, embedded as PNG.
    """
    import math
    import matplotlib
    matplotlib.use("Agg")
    import matplotlib.pyplot as plt
    import matplotlib.patches as mpatches
    import numpy as np

    # ── Shared helpers ────────────────────────────────────────────────────────
    logo_tmp = "/tmp/summit_logo_pdf.png"
    with open(logo_tmp, "wb") as f:
        f.write(logo_bytes)

    overall  = weighted_overall(scores)
    band     = score_band(overall)
    bg_rgb   = {"red":(253,233,229), "amber":(255,244,224), "green":(230,245,236)}
    fg_rgb2  = {"red":FG_RED_RGB,    "amber":FG_AMBER_RGB,   "green":FG_GREEN_RGB}
    fg_hex   = {"red":FG_RED,        "amber":FG_AMBER,        "green":FG_GREEN}
    lm_text  = {"red":"Needs Attention", "amber":"Developing", "green":"Good"}
    dim_short = {
        "Crawlability & Bot Access":"Crawlability",
        "Structured Data / Schema":"Structured Data",
        "LLM Content Signals":"LLM Content",
        "Meta & SEO Signals":"Meta / SEO",
        "Heading Structure":"Headings",
        "ARIA Implementation":"ARIA",
        "Link Quality":"Link Quality",
        "Image Alt Text":"Image Alt",
        "AI Search Health":"AI Health",
        "Duplicate Content & Tags":"Duplicates",
    }

    def hex_to_rgb01(h):
        h = h.lstrip("#")
        return tuple(int(h[i:i+2],16)/255 for i in (0,2,4))

    S_RED_01      = hex_to_rgb01(S_RED)
    S_CHARCOAL_01 = hex_to_rgb01(S_CHARCOAL)
    DIM_COLORS = {
        "Crawlability & Bot Access":   S_CHARCOAL_01,
        "Structured Data / Schema":    hex_to_rgb01("#7C3AED"),
        "LLM Content Signals":         S_RED_01,
        "Meta & SEO Signals":          hex_to_rgb01("#0E7C4A"),
        "Heading Structure":           hex_to_rgb01("#0369A1"),
        "ARIA Implementation":         hex_to_rgb01("#BE185D"),
        "Link Quality":                hex_to_rgb01("#78350F"),
        "Image Alt Text":              hex_to_rgb01("#374151"),
        "AI Search Health":            hex_to_rgb01("#B45309"),
        "Duplicate Content & Tags":    hex_to_rgb01("#6B6B80"),
    }

    sorted_scores = sorted(scores.items(), key=lambda x: x[1])

    # ── Build charts as PNG byte buffers ─────────────────────────────────────

    def make_bar_chart():
        dims  = [dim_short.get(d, d[:14]) for d, _ in sorted_scores]
        vals  = [v for _, v in sorted_scores]
        colors= [hex_to_rgb01(fg_hex[score_band(v)]) for v in vals]
        bg_cs = [hex_to_rgb01(
            {"red":BG_RED,"amber":BG_AMBER,"green":BG_GREEN}[score_band(v)]
        ) for v in vals]

        fig, ax = plt.subplots(figsize=(7.2, 3.8))
        fig.patch.set_facecolor("white")
        ax.set_facecolor("white")

        y = np.arange(len(dims))
        # Background track
        ax.barh(y, [10]*len(dims), height=0.6, color="#F3F3F5", zorder=1)
        # Score bar
        bars = ax.barh(y, vals, height=0.6, color=colors, zorder=2)
        # Score label
        for i, (bar, v) in enumerate(zip(bars, vals)):
            ax.text(v + 0.15, i, f"{v}/10",
                    va="center", ha="left", fontsize=8,
                    color=colors[i], fontweight="bold")
        # Threshold lines
        ax.axvline(3.3, color=hex_to_rgb01(FG_RED),   lw=0.8, ls="--", alpha=0.5, zorder=3)
        ax.axvline(6.6, color=hex_to_rgb01(FG_GREEN), lw=0.8, ls="--", alpha=0.5, zorder=3)
        ax.set_yticks(y)
        ax.set_yticklabels(dims, fontsize=8)
        ax.set_xlim(0, 12)
        ax.set_xticks([0,2,4,6,8,10])
        ax.tick_params(axis="x", labelsize=7)
        ax.set_xlabel("Score / 10", fontsize=8)
        ax.invert_yaxis()
        for spine in ax.spines.values(): spine.set_visible(False)
        ax.xaxis.grid(True, color="#EEEEEE", zorder=0)
        plt.tight_layout(pad=0.5)
        buf = io.BytesIO(); fig.savefig(buf, format="png", dpi=150, bbox_inches="tight"); buf.seek(0)
        plt.close(fig)
        return buf

    def make_radar_chart():
        dims = list(DIMENSION_CONFIG.keys())
        labels = [dim_short.get(d, d) for d in dims]
        vals   = [scores.get(d, 0) for d in dims]
        N = len(dims)
        angles = [n / float(N) * 2 * math.pi for n in range(N)]
        angles += angles[:1]; vals_plot = vals + vals[:1]

        fig, ax = plt.subplots(figsize=(4.5, 4.5), subplot_kw=dict(polar=True))
        fig.patch.set_facecolor("white")
        ax.set_facecolor("#FAFAFA")
        ax.plot(angles, vals_plot, color=S_RED_01, lw=2.0, zorder=3)
        ax.fill(angles, vals_plot, color=S_RED_01, alpha=0.12, zorder=2)
        ax.scatter(angles[:-1], vals, color=S_RED_01, s=25, zorder=4)
        ax.set_xticks(angles[:-1])
        ax.set_xticklabels(labels, size=6.5, color="#333333")
        ax.set_ylim(0, 10)
        ax.set_yticks([2, 4, 6, 8, 10])
        ax.set_yticklabels(["2","4","6","8","10"], size=6, color="#AAAAAA")
        ax.grid(color="#DDDDDD", linewidth=0.5)
        ax.spines["polar"].set_color("#CCCCCC")
        # Threshold circles
        theta = np.linspace(0, 2*math.pi, 100)
        ax.plot(theta, [3.3]*100, color=hex_to_rgb01(FG_RED),   lw=0.6, ls=":", alpha=0.5)
        ax.plot(theta, [6.6]*100, color=hex_to_rgb01(FG_GREEN), lw=0.6, ls=":", alpha=0.5)
        plt.tight_layout(pad=0.3)
        buf = io.BytesIO(); fig.savefig(buf, format="png", dpi=150, bbox_inches="tight"); buf.seek(0)
        plt.close(fig)
        return buf

    def make_score_tiles():
        """5×2 grid of score tiles."""
        items = sorted_scores  # 10 items
        ncols, nrows = 5, 2
        fig, axes = plt.subplots(nrows, ncols, figsize=(7.2, 2.0))
        fig.patch.set_facecolor("white")
        for idx, (dim, s) in enumerate(items):
            r, c = divmod(idx, ncols)
            ax = axes[r][c]
            bd = score_band(s)
            bg = tuple(x/255 for x in bg_rgb[bd])
            fg = tuple(x/255 for x in fg_rgb2[bd])
            ax.set_facecolor(bg)
            ax.set_xlim(0,1); ax.set_ylim(0,1)
            ax.axis("off")
            ax.text(0.5, 0.62, f"{s}/10", ha="center", va="center",
                    fontsize=14, fontweight="bold", color=fg,
                    transform=ax.transAxes)
            short = dim_short.get(dim, dim[:12])
            ax.text(0.5, 0.22, short, ha="center", va="center",
                    fontsize=6.5, color="#444444",
                    transform=ax.transAxes)
            for spine in ax.spines.values():
                spine.set_edgecolor("#FFFFFF"); spine.set_linewidth(2)
        plt.subplots_adjust(wspace=0.04, hspace=0.04)
        buf = io.BytesIO(); fig.savefig(buf, format="png", dpi=150, bbox_inches="tight",
                                        facecolor="white"); buf.seek(0)
        plt.close(fig)
        return buf

    # ── FPDF document ─────────────────────────────────────────────────────────
    class SummitPDF(FPDF):
        def header(self):
            self.image(logo_tmp, 12, 8, 20)
            self.set_font("Helvetica","B",14)
            self.set_text_color(*S_CHARCOAL_RGB)
            self.set_xy(35, 10)
            self.cell(0, 6, "AI Visibility Audit", ln=False)
            self.set_font("Helvetica","",8)
            self.set_text_color(130,130,140)
            self.set_xy(35, 17)
            self.cell(0, 5, _pdf_safe(f"{url}  -  {datetime.now().strftime('%B %Y')}", 110), ln=True)
            self.set_draw_color(*S_RED_RGB)
            self.set_line_width(0.8)
            self.line(12, 26, 198, 26)
            self.ln(2)

        def footer(self):
            self.set_y(-11)
            self.set_font("Helvetica","I",7.5)
            self.set_text_color(160,160,170)
            self.cell(0, 5,
                f"Summit AI Visibility Audit  -  Confidential  -  Page {self.page_no()}",
                align="C")

    pdf = SummitPDF()
    pdf.set_auto_page_break(auto=True, margin=14)

    # ════════════════════════════════════════════════════════════════════════
    # PAGE 1 — Hero + Score tiles + Bar chart
    # ════════════════════════════════════════════════════════════════════════
    pdf.add_page()

    # Hero block
    y0 = pdf.get_y()
    pdf.set_fill_color(*S_CHARCOAL_RGB)
    pdf.rect(12, y0, 52, 28, "F")
    pdf.set_text_color(255,255,255)
    pdf.set_font("Helvetica","B",26)
    pdf.set_xy(12, y0+2)
    pdf.cell(52, 12, f"{overall}/10", align="C", ln=False)
    pdf.set_font("Helvetica","",7)
    pdf.set_text_color(180,180,205)
    pdf.set_xy(12, y0+16)
    pdf.cell(52, 6, "Overall AI Visibility Score", align="C", ln=False)

    pdf.set_fill_color(*bg_rgb[band])
    pdf.rect(66, y0, 50, 28, "F")
    pdf.set_text_color(*fg_rgb2[band])
    pdf.set_font("Helvetica","B",13)
    pdf.set_xy(66, y0+7)
    pdf.cell(50, 8, lm_text[band], align="C", ln=False)
    pdf.set_font("Helvetica","",7)
    pdf.set_xy(66, y0+17)
    pdf.cell(50, 5, f"Weighted across {len(scores)} dimensions", align="C", ln=False)

    # Stats boxes
    n_crit = sum(1 for s in scores.values() if score_band(s)=="red")
    n_warn = sum(1 for s in scores.values() if score_band(s)=="amber")
    n_good = sum(1 for s in scores.values() if score_band(s)=="green")
    stat_items = [
        (n_crit,"Critical",  FG_RED_RGB,   (253,233,229)),
        (n_warn,"Warnings",  FG_AMBER_RGB, (255,244,224)),
        (n_good,"Passing",   FG_GREEN_RGB, (230,245,236)),
    ]
    for i,(val,lbl,col,bg) in enumerate(stat_items):
        bx = 119 + i*27
        pdf.set_fill_color(*bg)
        pdf.rect(bx, y0, 24, 28, "F")
        pdf.set_text_color(*col)
        pdf.set_font("Helvetica","B",17)
        pdf.set_xy(bx, y0+4)
        pdf.cell(24, 10, str(val), align="C", ln=False)
        pdf.set_font("Helvetica","",6.5)
        pdf.set_text_color(100,100,110)
        pdf.set_xy(bx, y0+17)
        pdf.cell(24, 5, lbl, align="C", ln=False)

    pdf.set_y(y0+32)
    pdf.ln(2)

    # Section heading helper
    def pdf_section(title):
        pdf.set_font("Helvetica","B",8.5)
        pdf.set_text_color(*S_CHARCOAL_RGB)
        pdf.cell(0, 5, title, ln=True)
        pdf.set_draw_color(*S_RED_RGB)
        pdf.set_line_width(0.5)
        pdf.line(12, pdf.get_y(), 198, pdf.get_y())
        pdf.ln(2)

    # Score tiles
    pdf_section("Dimension Scores")
    tiles_buf = make_score_tiles()
    tiles_tmp = "/tmp/pdf_tiles.png"
    with open(tiles_tmp,"wb") as f: f.write(tiles_buf.getvalue())
    pdf.image(tiles_tmp, x=12, w=186)
    pdf.ln(3)

    # Bar chart
    pdf_section("Score by Dimension")
    bar_buf = make_bar_chart()
    bar_tmp = "/tmp/pdf_bar.png"
    with open(bar_tmp,"wb") as f: f.write(bar_buf.getvalue())
    pdf.image(bar_tmp, x=12, w=186)

    # ════════════════════════════════════════════════════════════════════════
    # PAGE 2 — Radar + Issues table
    # ════════════════════════════════════════════════════════════════════════
    pdf.add_page()

    # Radar chart (left) + Score table (right)
    pdf_section("AI Visibility Profile")
    radar_buf = make_radar_chart()
    radar_tmp = "/tmp/pdf_radar.png"
    with open(radar_tmp,"wb") as f: f.write(radar_buf.getvalue())
    pdf.image(radar_tmp, x=12, w=90)

    # Score table alongside radar
    tx = 106; ty = pdf.get_y() - 80  # align with radar top
    pdf.set_xy(tx, ty)
    # Table header
    col_ws = [44, 12, 14, 18]
    hdrs   = ["Dimension","Wt","Score","Rating"]
    pdf.set_fill_color(*S_CHARCOAL_RGB)
    pdf.set_text_color(255,255,255)
    pdf.set_font("Helvetica","B",7)
    for w, h in zip(col_ws, hdrs):
        pdf.cell(w, 5, h, fill=True, border=0)
    pdf.ln()
    sev_lbl = {"red":"Needs Attn","amber":"Needs Work","green":"Good"}
    for dim, s in sorted_scores:
        bd = score_band(s)
        pdf.set_fill_color(250,250,252)
        pdf.set_text_color(50,50,60)
        pdf.set_font("Helvetica","",6.5)
        short = dim_short.get(dim, dim[:18])
        wp = int(DIMENSION_CONFIG.get(dim,{}).get("weight",0.05)*100)
        pdf.set_xy(tx, pdf.get_y())
        pdf.cell(col_ws[0], 5, f" {short}", fill=True, border=0)
        pdf.cell(col_ws[1], 5, f"{wp}%", fill=True, border=0, align="C")
        # Score cell coloured
        pdf.set_fill_color(*bg_rgb[bd])
        pdf.set_text_color(*fg_rgb2[bd])
        pdf.set_font("Helvetica","B",6.5)
        pdf.cell(col_ws[2], 5, f"{s}/10", fill=True, border=0, align="C")
        pdf.set_font("Helvetica","",6)
        pdf.cell(col_ws[3], 5, sev_lbl[bd], fill=True, border=0)
        pdf.ln()

    # Issues & Recommendations table
    pdf.set_y(pdf.get_y()+4)
    pdf_section("Issues & Recommendations")
    # Header
    iss_ws = [5, 17, 30, 146]
    pdf.set_fill_color(*S_CHARCOAL_RGB)
    pdf.set_text_color(255,255,255)
    pdf.set_font("Helvetica","B",7)
    for w, h in zip(iss_ws, ["#","Severity","Dimension","Recommendation"]):
        pdf.cell(w, 5, h, fill=True, border=0)
    pdf.ln()

    sev_col_map = {"critical":FG_RED_RGB,"warning":FG_AMBER_RGB,"info":(21,80,175)}
    sev_bg_map  = {"critical":(253,233,229),"warning":(255,244,224),"info":(235,240,251)}

    all_recs = []
    for d, res in all_results.items():
        w = DIMENSION_CONFIG.get(d,{}).get("weight",0.05)
        for iss in res.get("issues",[]):
            sev = iss.get("severity","info")
            all_recs.append({
                "pri": {"critical":1,"warning":2,"info":3}.get(sev,3),
                "w": w, "d": d,
                "rec": iss.get("recommendation",""),
                "sev": sev
            })
    all_recs.sort(key=lambda x:(x["pri"],-x["w"]))

    for rank, rec in enumerate(all_recs[:20], 1):
        sev = rec["sev"]
        fill_bg = (250,250,252) if rank%2==0 else (245,245,248)
        pdf.set_fill_color(*fill_bg)
        pdf.set_text_color(70,70,80)
        pdf.set_font("Helvetica","B",6.5)
        pdf.cell(iss_ws[0], 5, str(rank), fill=True, border=0, align="C")
        pdf.set_fill_color(*sev_bg_map.get(sev, fill_bg))
        pdf.set_text_color(*sev_col_map.get(sev,(80,80,80)))
        pdf.cell(iss_ws[1], 5, f" {sev.upper()}", fill=True, border=0)
        pdf.set_fill_color(*fill_bg)
        pdf.set_text_color(60,60,70)
        pdf.set_font("Helvetica","",6.5)
        pdf.cell(iss_ws[2], 5, f" {dim_short.get(rec['d'],rec['d'][:16])}", fill=True, border=0)
        rec_txt = _pdf_safe(rec["rec"], 180)
        pdf.cell(iss_ws[3], 5, f" {rec_txt}", fill=True, border=0)
        pdf.ln()

    # ════════════════════════════════════════════════════════════════════════
    # PAGE 3 — Executive Summary
    # ════════════════════════════════════════════════════════════════════════
    pdf.add_page()
    pdf_section("Executive Summary")
    paras = [p.strip() for p in exec_summary.split("\n\n") if p.strip()]
    pdf.set_font("Helvetica","",9)
    pdf.set_text_color(35,35,45)
    for para in paras:
        pdf.multi_cell(0, 5.2, _pdf_safe(para, 1200))
        pdf.ln(1.5)

    raw = pdf.output()
    return bytes(raw) if not isinstance(raw, bytes) else raw



def run_single_audit(client, url, progress_bar=None, status_box=None, manual_html=None):
    def prog(v):
        if progress_bar is not None:
            progress_bar.progress(v)
    def stat(t):
        if status_box is not None:
            status_box.text(t)

    stat("📡 Fetching page…"); prog(5)
    if manual_html:
        # Use manually pasted HTML — bypasses all bot protection
        page_data = {"url":url,"status_code":200,"error":None,"load_time":None,
                     "is_https":url.startswith("https://"),"redirect_chain":[],
                     "fetch_source":"manual_paste","is_js_shell":False}
        page_data.update(_extract_signals(url, manual_html))
        stat("📋 Using manually pasted HTML…"); prog(8)
    else:
        page_data = fetch_page(url)
    if page_data.get("error"):
        return None, None, None, page_data

    stat("🤖 Checking robots.txt & llms.txt…"); prog(10)
    page_data["robots_data"] = check_robots(url)
    page_data["llms_txt"]    = check_llms_txt(url)

    scores, all_results = {}, {}
    dims = list(DIMENSION_CONFIG.keys())
    for i, dim in enumerate(dims):
        stat(f"🔍 Analysing: {dim}…"); prog(15 + int(68 * i / len(dims)))
        res = analyse_dimension(client, page_data, dim)
        scores[dim] = res.get("score", 0)
        all_results[dim] = res
        time.sleep(0.2)

    stat("✍️ Generating executive summary…"); prog(88)
    exec_summary = gen_exec_summary(client, url, scores, all_results)
    prog(100)
    return scores, all_results, exec_summary, page_data


# =============================================================================
# DASHBOARD RENDERER — for a single audit result
# =============================================================================
def render_dashboard(url, scores, all_results, exec_summary, page_data,
                     logo_bytes, logo_b64_str, tab_key=""):
    overall = weighted_overall(scores)
    band    = score_band(overall)
    bg_map  = {"red":BG_RED,"amber":BG_AMBER,"green":BG_GREEN}
    fg_map  = {"red":FG_RED,"amber":FG_AMBER,"green":FG_GREEN}
    lbl_map = {"red":"Needs Attention","amber":"Developing","green":"Good"}

    # Hero row
    c1,c2,c3,c4 = st.columns([1.6,1.4,1,1])
    with c1:
        st.markdown(f"""
        <div class="score-hero">
            <div class="lbl">Overall AI Visibility Score</div>
            <div><span class="num">{overall}</span><span class="denom"> /10</span></div>
            <div style="margin-top:8px;background:rgba(255,255,255,0.12);border-radius:7px;
                        padding:5px 16px;display:inline-block;font-weight:700;font-size:0.96rem">
                {lbl_map[band]}
            </div>
        </div>
        """, unsafe_allow_html=True)
    with c2:
        st.plotly_chart(gauge_chart(overall), use_container_width=True, config={"displayModeBar":False})
    with c3:
        n_crit=sum(1 for r in all_results.values() for iss in r.get("issues",[]) if iss.get("severity")=="critical")
        n_warn=sum(1 for r in all_results.values() for iss in r.get("issues",[]) if iss.get("severity")=="warning")
        st.markdown(f"""
        <div class="mcard red" style="text-align:center"><div class="mnum">{n_crit}</div><div class="mlbl">Critical Issues</div></div>
        <div class="mcard amber" style="text-align:center"><div class="mnum">{n_warn}</div><div class="mlbl">Warnings</div></div>
        """, unsafe_allow_html=True)
    with c4:
        n_pass=sum(1 for s in scores.values() if s>6.5)
        lt=page_data.get("load_time","—")
        st.markdown(f"""
        <div class="mcard green" style="text-align:center"><div class="mnum">{n_pass}</div><div class="mlbl">Passing</div></div>
        <div class="mcard grey" style="text-align:center"><div class="mnum">{lt}s</div><div class="mlbl">Load Time</div></div>
        """, unsafe_allow_html=True)

    st.markdown("---")
    t1,t2,t3,t4,t5 = st.tabs(["📊 Scores","🕸️ Radar","⚖️ Weighting","📋 Issues","📝 Summary"])

    with t1:
        ba,bb=st.columns([2,1])
        with ba:
            st.markdown("#### Score by Dimension")
            st.plotly_chart(bar_chart(scores), use_container_width=True, config={"displayModeBar":False})
        with bb:
            st.markdown("#### Breakdown")
            rows=""
            for d,s in sorted(scores.items(),key=lambda x:x[1]):
                bc=score_band(s); ic=DIMENSION_CONFIG.get(d,{}).get("icon","•")
                rows+=f'<tr><td>{ic} {d}</td><td style="text-align:right"><span class="badge-{bc}">{s}/10</span></td></tr>'
            st.markdown(f'<table class="at"><thead><tr><th>Dimension</th><th>Score</th></tr></thead><tbody>{rows}</tbody></table>',unsafe_allow_html=True)

    with t2:
        st.plotly_chart(radar_chart(scores), use_container_width=True, config={"displayModeBar":False})

    with t3:
        wa,wb=st.columns(2)
        with wa:
            st.plotly_chart(weight_donut(scores), use_container_width=True, config={"displayModeBar":False})
        with wb:
            rows=""
            for d,cfg in DIMENSION_CONFIG.items():
                s=scores.get(d,0); bc=score_band(s); ws=round(s*cfg["weight"],2)
                rows+=f'<tr><td>{cfg["icon"]} {d}</td><td style="text-align:center">{int(cfg["weight"]*100)}%</td><td style="text-align:center"><span class="badge-{bc}">{s}/10</span></td><td style="text-align:center;font-weight:600;color:{S_CHARCOAL}">{ws}</td></tr>'
            st.markdown(f'<table class="at"><thead><tr><th>Dimension</th><th>Weight</th><th>Score</th><th>Weighted</th></tr></thead><tbody>{rows}</tbody></table>',unsafe_allow_html=True)

    with t4:
        sf=st.selectbox("Filter",["All","Critical","Warning","Info"],key=f"sf_{tab_key}")
        flat=sorted([{"sev":i.get("severity","info"),"dim":d,**i} for d,r in all_results.items() for i in r.get("issues",[])],
                    key=lambda x:{"critical":0,"warning":1,"info":2}.get(x["sev"],3))
        shown=0
        for iss in flat:
            if sf!="All" and iss["sev"].lower()!=sf.lower(): continue
            shown+=1
            ic={"critical":"🔴","warning":"🟡","info":"🔵"}.get(iss["sev"],"⚪")
            st.markdown(f"""
            <div class="irow {iss['sev']}">
                <div style="min-width:22px">{ic}</div>
                <div style="flex:1">
                    <div><span class="ititle">{iss.get('issue','')}</span><span class="idim">{iss['dim']}</span></div>
                    <div class="irec">💡 {iss.get('recommendation','')}</div>
                </div>
            </div>""", unsafe_allow_html=True)
        if shown==0: st.info("No issues for this filter.")

    with t5:
        st.markdown(f"""
        <div style="background:{S_WHITE};border-radius:11px;padding:24px;box-shadow:0 2px 7px rgba(0,0,0,0.07);
        color:{S_INK};line-height:1.7;font-size:0.92rem">{exec_summary.replace(chr(10),"<br>")}</div>
        """, unsafe_allow_html=True)

    st.markdown("---")
    st.markdown("### 🔎 Findings by Dimension")
    for d,res in all_results.items():
        s=scores.get(d,0); bd=score_band(s); ic=DIMENSION_CONFIG.get(d,{}).get("icon","•")
        with st.expander(f"{ic} {d}  —  {s}/10", expanded=(s<=3)):
            da,db=st.columns([1,2])
            with da:
                st.markdown(f"""
                <div style="text-align:center;padding:16px;background:{bg_map[bd]};border-radius:9px;border:1px solid {S_LIGHT}">
                    <div style="font-size:2.6rem;font-weight:800;color:{fg_map[bd]}">{s}</div>
                    <div style="color:{fg_map[bd]};font-size:0.74rem;font-weight:600;text-transform:uppercase;letter-spacing:0.05em">/10 — {lbl_map[bd]}</div>
                </div>
                <p style="text-align:center;color:{S_CAPTION};font-size:0.76rem;margin-top:6px">Weight: {int(DIMENSION_CONFIG.get(d,{}).get('weight',0.05)*100)}%</p>
                """, unsafe_allow_html=True)
            with db:
                if res.get("summary"): st.markdown(f"**{res['summary']}**")
                for ft in res.get("findings",[]): st.markdown(f"- {ft}")
            issues=res.get("issues",[])
            if issues:
                st.markdown("**Issues**")
                for iss in issues:
                    sev=iss.get("severity","info")
                    ic2={"critical":"🔴","warning":"🟡","info":"🔵"}.get(sev,"⚪")
                    st.markdown(f"{ic2} **{iss.get('issue','')}**")
                    st.caption(f"Recommendation: {iss.get('recommendation','')}")
            pos=res.get("positive",[])
            if pos:
                st.markdown("**Strengths**")
                for p in pos: st.markdown(f"✅ {p}")

    # Exports — stored in session_state to survive downloads
    sk = tab_key or "single"
    st.markdown("---")
    st.markdown("### 📥 Export Reports")
    e1,e2,e3=st.columns(3)
    domain=urlparse(url).netloc.replace("www.",""); stamp=datetime.now().strftime("%Y%m")

    # Pre-build exports into session_state so downloads don't trigger full re-run
    docx_key=f"docx_{sk}"; pdf_key=f"pdf_{sk}"; json_key=f"json_{sk}"
    if docx_key not in st.session_state:
        with st.spinner("Building .docx…"):
            st.session_state[docx_key] = build_docx(url,scores,all_results,exec_summary,logo_bytes)
    if pdf_key not in st.session_state:
        with st.spinner("Building PDF…"):
            st.session_state[pdf_key] = build_pdf_with_issues(url,scores,all_results,exec_summary,logo_bytes)
    if json_key not in st.session_state:
        export={"url":url,"audit_date":datetime.now().isoformat(),"overall_score":overall,
                "scores":scores,"results":{k:{kk:vv for kk,vv in v.items() if kk!="soup"} for k,v in all_results.items()},
                "executive_summary":exec_summary}
        st.session_state[json_key] = json.dumps(export,indent=2,default=str)

    with e1:
        st.markdown(f"""<div style="background:{S_WHITE};border-radius:9px;padding:16px;box-shadow:0 2px 6px rgba(0,0,0,0.06);border-top:4px solid {S_CHARCOAL}">
            <h4 style="color:{S_CHARCOAL};margin-top:0">📄 Word Report</h4>
            <p style="color:{S_MUTED};font-size:0.81rem">Full Summit-branded report with findings and recommendations.</p></div>""",unsafe_allow_html=True)
        st.download_button("⬇️ Download .docx", st.session_state[docx_key],
            file_name=f"summit_ai_audit_{domain}_{stamp}.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            use_container_width=True, key=f"dl_docx_{sk}")

    with e2:
        st.markdown(f"""<div style="background:{S_WHITE};border-radius:9px;padding:16px;box-shadow:0 2px 6px rgba(0,0,0,0.06);border-top:4px solid {S_RED}">
            <h4 style="color:{S_RED};margin-top:0">📋 PDF One-Pager</h4>
            <p style="color:{S_MUTED};font-size:0.81rem">Visual dashboard summary — score tiles, priority actions, exec summary.</p></div>""",unsafe_allow_html=True)
        st.download_button("⬇️ Download PDF", st.session_state[pdf_key],
            file_name=f"summit_ai_audit_{domain}_{stamp}.pdf",
            mime="application/pdf",
            use_container_width=True, key=f"dl_pdf_{sk}")

    with e3:
        st.markdown(f"""<div style="background:{S_WHITE};border-radius:9px;padding:16px;box-shadow:0 2px 6px rgba(0,0,0,0.06);border-top:4px solid {FG_GREEN}">
            <h4 style="color:{FG_GREEN};margin-top:0">📊 JSON Data</h4>
            <p style="color:{S_MUTED};font-size:0.81rem">Raw audit data for integration with other tools or CRMs.</p></div>""",unsafe_allow_html=True)
        st.download_button("⬇️ Download JSON", st.session_state[json_key],
            file_name=f"summit_ai_audit_{domain}_{stamp}.json",
            mime="application/json",
            use_container_width=True, key=f"dl_json_{sk}")


# =============================================================================
# MAIN
# =============================================================================
def main():
    inject_css()
    logo_bytes   = load_logo_bytes()
    logo_b64_str = logo_b64(logo_bytes)

    # Session state init
    if "audit_results" not in st.session_state:
        st.session_state.audit_results = {}   # url -> {scores, all_results, exec_summary, page_data}

    # ── Sidebar ──────────────────────────────────────────────────────────
    with st.sidebar:
        st.markdown(f'<img src="data:image/png;base64,{logo_b64_str}" width="90" style="margin-bottom:18px;display:block">',
                    unsafe_allow_html=True)
        st.markdown("## AI Visibility Audit")
        st.markdown("---")

        # API key
        api_key=""
        try:
            api_key=st.secrets["GEMINI_API_KEY"]
        except Exception:
            api_key=st.text_input("Gemini API Key",type="password",
                                  help="Or add GEMINI_API_KEY to .streamlit/secrets.toml")
        st.markdown("---")

        # URL input mode toggle
        mode=st.radio("Audit mode",["Single URL","Multiple URLs"],horizontal=True)

        if mode=="Single URL":
            url_single=st.text_input("URL to Audit",placeholder="https://example.com")
            urls_to_audit=[url_single] if url_single.strip() else []
        else:
            url_text=st.text_area("URLs to Audit (one per line)",
                                   placeholder="https://example.com\nhttps://example.com/category",
                                   height=110,
                                   help="Enter up to 5 URLs, one per line")
            urls_to_audit=[u.strip() for u in url_text.strip().splitlines() if u.strip()][:5]
            if len(urls_to_audit)>0:
                st.caption(f"{len(urls_to_audit)} URL{'s' if len(urls_to_audit)!=1 else ''} queued")

        # Manual HTML paste — for sites that block automated fetching
        with st.expander("📋 Paste HTML manually (optional)"):
            st.caption("If the site blocks automated fetching, paste its raw HTML here. Right-click the page in your browser → View Page Source → Select All → paste here.")
            st.caption("Single URL: applies to that URL. Multiple URLs: applies to the first URL only. Other URLs are fetched automatically.")
            st.text_area("Paste raw HTML", height=120, placeholder="<!DOCTYPE html>...", key="pasted_html_input")

        st.markdown("---")
        run_btn=st.button("🚀 Run Audit", use_container_width=True)
        if st.button("🗑️ Clear Results", use_container_width=True):
            st.session_state.audit_results={}
            # Clear cached exports
            for k in list(st.session_state.keys()):
                if k.startswith(("docx_","pdf_","json_")): del st.session_state[k]
            st.rerun()

        st.markdown("---")
        st.caption("Summit AI Visibility Audit v2.1")
        st.caption("Powered by Gemini 2.5 Flash")
        st.caption("© Summit Performance Marketing")

    # ── Header ───────────────────────────────────────────────────────────
    hc1,hc2=st.columns([3,1])
    with hc1:
        st.markdown(
            f'<div style="display:flex;align-items:center;gap:14px;margin-bottom:4px">'
            f'<img src="data:image/png;base64,{logo_b64_str}" width="58" style="border-radius:6px">'
            f'<div><h1 style="margin:0;color:{S_CHARCOAL};font-size:1.7rem;font-weight:800">AI Visibility Audit</h1>'
            f'<p style="margin:0;color:{S_MUTED};font-size:0.86rem">Technical audit for AI crawler access &amp; content extraction quality</p>'
            f'</div></div>', unsafe_allow_html=True)
    with hc2:
        st.markdown(f'<p style="text-align:right;color:{S_CAPTION};margin-top:20px;font-size:0.8rem">{datetime.now().strftime("%d %b %Y")}</p>',unsafe_allow_html=True)
    st.markdown("---")

    # ── Run audit if button pressed ───────────────────────────────────────
    if run_btn:
        if not api_key:
            st.error("Please enter your Gemini API key."); return
        if not urls_to_audit:
            st.error("Please enter at least one URL."); return

        try:
            client=genai.Client(api_key=api_key)
        except Exception as e:
            st.error(f"Gemini API error: {e}"); return

        for url_raw in urls_to_audit:
            url=url_raw if url_raw.startswith(("http://","https://")) else "https://"+url_raw
            st.markdown(f"#### Auditing `{url}`")
            prog=st.progress(0); stat=st.empty()

            is_first_url = (url_raw == urls_to_audit[0])
            manual_html = st.session_state.get("pasted_html_input","").strip() if is_first_url else ""
            scores,all_results,exec_summary,page_data = run_single_audit(
                client, url,
                progress_bar=prog,
                status_box=stat,
                manual_html=manual_html if manual_html else None,
            )
            prog.empty(); stat.empty()

            if scores is None:
                st.error(f"Failed to fetch {url}: {page_data.get('error','')}"); continue

            # Cache in session state
            st.session_state.audit_results[url] = {
                "scores":scores,"all_results":all_results,
                "exec_summary":exec_summary,"page_data":page_data,
            }
            # Clear any stale exports for this URL
            slug=url.replace("https://","").replace("http://","").replace("/","_")[:40]
            for k in list(st.session_state.keys()):
                if k.endswith(slug): del st.session_state[k]

        st.rerun()

    # ── Render results ────────────────────────────────────────────────────
    results = st.session_state.audit_results

    if not results:
        # Landing
        st.markdown(f"""
        <div style="background:{S_WHITE};border-radius:12px;padding:34px;text-align:center;
        box-shadow:0 2px 8px rgba(0,0,0,0.07);margin-bottom:26px">
            <h2 style="color:{S_CHARCOAL};margin-top:0">Enter a URL in the sidebar to begin</h2>
            <p style="color:{S_MUTED};max-width:460px;margin:10px auto 0">Audits 10 AI visibility dimensions weighted by
            impact. Generates a visual dashboard plus branded Word and PDF exports.</p>
        </div>""", unsafe_allow_html=True)
        st.markdown("### Audit Dimensions")
        cols=st.columns(2)
        for i,(d,cfg) in enumerate(DIMENSION_CONFIG.items()):
            with cols[i%2]:
                st.markdown(f"""
                <div style="background:{S_WHITE};border-radius:8px;padding:12px 15px;margin-bottom:8px;
                box-shadow:0 1px 4px rgba(0,0,0,0.05);border-left:4px solid {cfg['color']}">
                    <div style="display:flex;justify-content:space-between;align-items:center">
                        <span style="font-weight:600;font-size:0.88rem;color:{S_CHARCOAL}">{cfg['icon']} {d}</span>
                        <span style="background:{S_OFFWHITE};color:{S_CHARCOAL};padding:1px 7px;border-radius:9px;font-size:0.74rem;font-weight:700">{int(cfg['weight']*100)}%</span>
                    </div>
                    <p style="margin:3px 0 0;color:{S_MUTED};font-size:0.78rem">{cfg['description']}</p>
                </div>""", unsafe_allow_html=True)
        return

    # Single result — clean single view
    if len(results)==1:
        url,data = list(results.items())[0]
        slug=url.replace("https://","").replace("http://","").replace("/","_")[:40]
        render_dashboard(url,data["scores"],data["all_results"],data["exec_summary"],
                         data["page_data"],logo_bytes,logo_b64_str,tab_key=slug)
        return

    # Multiple results — tabbed by URL + comparison view
    urls=list(results.keys())
    short_labels={u: urlparse(u).netloc.replace("www.","")+urlparse(u).path[:24] for u in urls}
    tab_labels=["📊 Compare"]+[f"🔗 {short_labels[u]}" for u in urls]
    tabs=st.tabs(tab_labels)

    with tabs[0]:
        st.markdown("### Side-by-side Comparison")
        # Overall scores row
        score_cols=st.columns(len(urls))
        for i,url in enumerate(urls):
            s=results[url]["scores"]; ov=weighted_overall(s); bd=score_band(ov)
            bg={"red":BG_RED,"amber":BG_AMBER,"green":BG_GREEN}[bd]
            fg={"red":FG_RED,"amber":FG_AMBER,"green":FG_GREEN}[bd]
            with score_cols[i]:
                st.markdown(f"""
                <div style="background:{bg};border-radius:10px;padding:14px;text-align:center;border-top:4px solid {fg}">
                    <div style="font-size:0.72rem;font-weight:600;color:{S_MUTED};text-transform:uppercase;letter-spacing:0.06em;margin-bottom:4px;overflow:hidden;text-overflow:ellipsis;white-space:nowrap"
                         title="{url}">{short_labels[url]}</div>
                    <div style="font-size:2.4rem;font-weight:800;color:{fg};line-height:1">{ov}</div>
                    <div style="font-size:0.75rem;color:{fg}">/10</div>
                </div>""", unsafe_allow_html=True)

        st.markdown("---")
        cc1,cc2=st.columns(2)
        with cc1:
            st.markdown("#### Radar Comparison")
            scores_map={short_labels[u]:results[u]["scores"] for u in urls}
            st.plotly_chart(multi_radar_chart(scores_map),use_container_width=True,config={"displayModeBar":False})
        with cc2:
            st.markdown("#### Score Comparison")
            st.plotly_chart(multi_bar_chart(scores_map),use_container_width=True,config={"displayModeBar":False})

        st.markdown("#### Dimension Comparison Table")
        dims=list(DIMENSION_CONFIG.keys())
        hdr="<tr><th>Dimension</th>"+"".join(f"<th>{short_labels[u]}</th>" for u in urls)+"</tr>"
        body=""
        for d in dims:
            row=f"<tr><td>{DIMENSION_CONFIG[d]['icon']} {d}</td>"
            for u in urls:
                s=results[u]["scores"].get(d,0); bc=score_band(s)
                row+=f'<td style="text-align:center"><span class="badge-{bc}">{s}/10</span></td>'
            body+=row+"</tr>"
        st.markdown(f'<table class="at"><thead>{hdr}</thead><tbody>{body}</tbody></table>',unsafe_allow_html=True)

    for i,url in enumerate(urls):
        with tabs[i+1]:
            data=results[url]
            slug=url.replace("https://","").replace("http://","").replace("/","_")[:40]
            render_dashboard(url,data["scores"],data["all_results"],data["exec_summary"],
                             data["page_data"],logo_bytes,logo_b64_str,tab_key=slug)


if __name__=="__main__":
    main()
